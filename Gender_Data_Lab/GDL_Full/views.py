# Django Imports
from django.shortcuts import render, redirect
from django.contrib.auth import authenticate, login, logout as django_logout
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.core.exceptions import ValidationError
from django.core.files.storage import FileSystemStorage
from django.conf import settings
from django.http import HttpResponse, JsonResponse
from django.views.decorators.csrf import csrf_exempt
from .models import *
# Python/Third-party Imports
import os
import pandas as pd
import base64
from datetime import datetime
import json
from io import BytesIO
from docx import Document
from docx.shared import Inches
import openai
import matplotlib.pyplot as plt 
# Project Imports
from .models import DataFileUpload
from  .import indicator_functions
#=====CHAT BOT INMPORTS=================
from django.shortcuts import render
from django.views.decorators.csrf import csrf_exempt
from django.http import JsonResponse
from datetime import datetime
import openai
import pandas as pd
import json


# ===== Constants =====
path = 'C:/Users/DBORINGO/Desktop/Miscelenous/GDL/Gender_Data_Lab/Gender_Data_Lab/survey_files/survey_files/eice5.csv'
UPLOAD_DIR = os.path.join(settings.MEDIA_ROOT)
REPORTS_DIR = os.path.join(settings.MEDIA_ROOT, 'reports')

# ===== Gender_Data_Lab_App Views =====
def index(request):
    return render(request, 'GDL_Full/index.html')

def instruction(request):
    return render(request, 'GDL_Full/instruction.html')


# ===== Warehouse_NISR_App Views =====
def admin_login(request):
    if request.method == "POST":
        username = request.POST.get("username")
        password = request.POST.get("password")
        user = authenticate(request, username=username, password=password)
        if user:
            login(request, user)
            next_page = request.GET.get('next', '/dashboard/')
            return redirect(next_page)
        else:
            messages.error(request, "Invalid username or password.")
    return render(request, 'GDL_Full/Warehouse_Login.html')

@login_required
def dashboard(request):
    uploaded_files = DataFileUpload.objects.all().order_by('-uploaded_at')
    return render(request, 'GDL_Full/dashboard.html', {'uploaded_files': uploaded_files})

@login_required
def logout(request):
    django_logout(request)
    return redirect('login')

@login_required
def fileupload(request):
    return render(request, 'GDL_Full/fileupload.html')

@login_required
def upload_file(request):
    if request.method == 'POST':
        uploaded_file = request.FILES.get('dataFile')
        source_type = request.POST.get('source_type')
        start_date = request.POST.get('start_date') or None
        end_date = request.POST.get('end_date') or None
        time_frequency = request.POST.get('time_frequency') or None
        survey_name = request.POST.get('survey_name') or None
        internal_description = request.POST.get('internal_description') or None
        external_file_name = request.POST.get('external_file_name') or None
        file_section_from = request.POST.get('file_section_from') or None
        external_description = request.POST.get('external_description') or None

        if not uploaded_file:
            messages.error(request, "No file selected.")
            return redirect('fileupload')

        try:
            data_file = DataFileUpload(
                file=uploaded_file,
                source_type=source_type,
                start_date=start_date,
                end_date=end_date,
                time_frequency=time_frequency,
                survey_name=survey_name,
                internal_description=internal_description,
                external_file_name=external_file_name,
                file_section_from=file_section_from,
                external_description=external_description,
            )
            data_file.save()
            messages.success(request, f"File '{uploaded_file.name}' uploaded successfully.")
        except ValidationError as e:
            messages.error(request, f"Upload failed: {'; '.join(e.messages)}")

    return redirect('fileupload')


# ===== Indicator Search Views =====
def indicator_search(request):
    surveys = Survey.objects.all()
    return render(request, 'GDL_Full/indicator_search.html', {'surveys': surveys})

def upload_indicator_file(request):
    if request.method == 'POST' and request.FILES.get('upload_file'):
        upload = request.FILES['upload_file']
        ext = os.path.splitext(upload.name)[1].lower()
        if ext not in ['.csv', '.xlsx', '.sav', '.dta']:
            return HttpResponse("Unsupported file type", status=400)

        fs = FileSystemStorage(location=UPLOAD_DIR)
        fs.save(upload.name, upload)

    return redirect('indicator_search')

def analyze_indicator(request, survey_id):
    if request.method == 'POST':
        survey = Survey.objects.get(survey_id=survey_id)
        dataset = survey.file.url
        indicator = request.POST.get('indicator').strip().lower()
        file_path = os.path.join(path, dataset)

        try:
            df = pd.read_csv(path) if dataset.endswith('.csv') else pd.read_excel(path)
        except Exception as e:
            return HttpResponse(f"Error reading file: {e}", status=500)

        if indicator == 'population structure (%), by sex and five-year age group':
              table = indicator_functions.population_structure_by_sex_five_year_group(df)
              fg,fig = indicator_functions.plot_population_structure_by_sex_five_year_group(df)
        elif indicator == 'sex of the household heads by province':
             table = indicator_functions.sex_of_household_heads_tbl(df)
             fg,fig = indicator_functions.sex_of_household_heads_plot(df)
        else:
            return HttpResponse("Unsupported indicator", status=400)

        # Convert plot to base64
        buffer = BytesIO()
        fg.savefig(buffer, format='png')
        buffer.seek(0)
        img_base64 = base64.b64encode(buffer.read()).decode('utf-8')
        buffer.close()
        plt.close(fg)

        surveys = Survey.objects.all()
    # return render(request, 'Gender_Data_Lab_App/indicator_search.html', {'surveys': surveys})
       
        return render(request, 'GDL_Full/indicator_search.html', {
            'surveys': surveys,
            'selected_dataset': path,
            'selected_indicator': indicator,
            'img_base64': img_base64,
            'table_html': table.to_html(classes="table-auto border", index=False),
            'show_download': True,
            'indicator': indicator.capitalize(),
        })

    return redirect('indicator_search')


# ====================== download report =====================
def download_report(request):
    if request.method == 'POST':
        dataset = request.POST.get('dataset')
        indicator = request.POST.get('indicator')
        file_path = os.path.join(UPLOAD_DIR, dataset)

        df = pd.read_csv(file_path) if dataset.endswith('.csv') else pd.read_excel(file_path)

    if indicator == 'population structure (%), by sex and five-year age group':
        table = indicator_functions.population_structure_by_sex_five_year_group(df)
        fg,fig = indicator_functions.plot_population_structure_by_sex_five_year_group(df)
    elif indicator == 'sex of the household heads by province':
        table = indicator_functions.sex_of_household_heads_tbl(df)
        fg,fig = indicator_functions.sex_of_household_heads_plot(df)
    else:
        return HttpResponse("Invalid indicator", status=400)

    # Create Word doc
    doc = Document()
    doc.add_heading('Indicator Report', 0)

    # Save figure to image
    image_stream = BytesIO()
    fg.savefig(image_stream, format='png')
    image_stream.seek(0)
    image_path = os.path.join(UPLOAD_DIR, 'temp_plot.png')
    with open(image_path, 'wb') as f:
        f.write(image_stream.read())
    plt.close(fg)

    doc.add_picture(image_path, width=Inches(6))
    os.remove(image_path)

    # Add table
    doc.add_heading('Summary Table', level=2)
    t = doc.add_table(rows=1, cols=len(table.columns))
    hdr_cells = t.rows[0].cells
    for i, col in enumerate(table.columns):
        hdr_cells[i].text = str(col)

    for _, row in table.iterrows():
        row_cells = t.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)

    doc_stream = BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)

    response = HttpResponse(doc_stream.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=indicator_report.docx'
    return response

    return redirect('indicator_search')

def get_indicators(request, survey_id):
    indicators = Indicator.objects.filter(survey__survey_id=survey_id)
    indicator_list = [{'name': ind.indicator_name} for ind in indicators]

    return JsonResponse({'indicators': indicator_list})


#===================AI ANALYTICS============================


# Load vector DB
df = pd.read_pickle('C:/Users/DBORINGO/Desktop/Miscelenous/GDL/Gender_Data_Lab/Gender_Data_Lab/minagri_report_vector_db_pickle.pkl')  # Adjust path as needed

# Set your API key
openai.api_key = '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'

# Import or define your chatbot logic here
from document_chat_agents import enhance_query, query_system, respond_with_retrieval, respond_without_retrieval, detect_and_translate_to_english, translate_to_user_language

def ai_analytics(request):
    return render(request, "GDL_Full/ai_analytics.html")

@csrf_exempt
def chat(request):
    if request.method == "POST":
        data = json.loads(request.body)
        user_input = data.get("message", "")
        if not user_input:
            return JsonResponse({"error": "Empty input"}, status=400)

        
        # 1️⃣ Detect language and translate user input
        translation_result = detect_and_translate_to_english(user_input)
        translated_input = translation_result["translated_text"]
        user_language = translation_result["detected_language"]

        # 2️⃣ Enhance and classify query
        enhanced_result = enhance_query(translated_input)
        enhanced_query = enhanced_result.get("enhanced_query", translated_input)
        needs_retrieval = enhanced_result.get("retrievals_required", True)

        # 3️⃣ Route to correct agent
        if needs_retrieval:
            retrieved_rows = query_system(enhanced_query, df)
            assistant_reply = respond_with_retrieval(enhanced_query, retrieved_rows)
        else:
            assistant_reply = respond_without_retrieval(enhanced_query)

        # 4️⃣ Translate assistant's reply back to user language (if needed)
        if user_language.lower() != "english":
            assistant_reply = translate_to_user_language(assistant_reply, user_language)


        return JsonResponse({
            "reply": assistant_reply,
            "timestamp": datetime.now().strftime("%H:%M")
        })




        